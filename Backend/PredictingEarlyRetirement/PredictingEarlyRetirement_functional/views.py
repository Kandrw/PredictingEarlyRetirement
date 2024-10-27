from django.shortcuts import render

import logging
import os
import time
import subprocess
from subprocess import Popen, PIPE
from PIL import Image
from pathlib import Path
import io
import PIL
import sys


from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework import status
from django.http import HttpResponse, JsonResponse
import base64
from django.core.files.base import ContentFile

from PredictingEarlyRetirement_functional.models import ImageModel
from django.core.files.uploadedfile import InMemoryUploadedFile

from django.conf import settings

from django.http import JsonResponse
from rest_framework.decorators import api_view
from django.core.files.storage import default_storage
from docx import Document

class ImageUploadView(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        csv_file = request.FILES.get('file')
        
        # Проверка на наличие файла и его расширение
        if not csv_file or not csv_file.name.endswith('.csv'):
            return Response({"error": "Пожалуйста, загрузите файл в формате .csv."}, status=status.HTTP_400_BAD_REQUEST)

        # Путь для сохранения загруженного файла
        upload_dir = os.path.join(settings.MEDIA_ROOT, 'files')
        os.makedirs(upload_dir, exist_ok=True)  # Создание директории, если она не существует

        file_path = os.path.join(upload_dir, csv_file.name)
        print(file_path)
        # Сохранение файла
        with open(file_path, 'wb+') as destination:
            for chunk in csv_file.chunks():
                destination.write(chunk)

        return Response({"message": "Файл успешно загружен!", "file_path": file_path}, status=status.HTTP_201_CREATED)

class FileListView(APIView):
    def get(self, request, *args, **kwargs):
        # Путь к директории с файлами
        directory = os.path.join(settings.MEDIA_ROOT, 'files')
        try:
            # Получение списка файлов в директории
            files = [f for f in os.listdir(directory) if f.endswith('.csv')]
            return Response({"files": files}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def submit_report(request):
    serial_number = request.data.get('serialNumber')
    defects = request.data.get('defects')

    if not defects:
        return JsonResponse({'error': 'Defects data is missing or empty'}, status=400)

    doc = Document()
    doc.add_heading('Отчет по дефектам', level=1)
    doc.add_paragraph(f'Серийный номер: {serial_number}')

    # for key, value in defects.items():
    #     doc.add_paragraph(f'{key}: {value}')
    for key, value in list(defects.items())[:-1]:
        doc.add_paragraph(f'{key}: {value}')
    file_path = f'reports/report_{serial_number}.docx'
    doc.save(file_path)

    with open(file_path, 'rb') as f:
        response = JsonResponse({'status': 'success'})
        response['Content-Disposition'] = f'attachment; filename="report_{serial_number}.docx"'
        response.write(f.read())
        return response

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

@api_view(['POST'])
def submit_report_pdf(request):
    serial_number = request.data.get('serialNumber')
    defects = request.data.get('defects')

    if not defects:
        return JsonResponse({'error': 'Defects data is missing or empty'}, status=400)

    file_path = f'reports/report_{serial_number}.pdf'
    c = canvas.Canvas(file_path, pagesize=letter)
    c.drawString(100, 750, f'Report about Laptop defects')
    c.drawString(100, 735, f'Serial Number: {serial_number}')

    y = 700
    # for key, value in defects.items():
    #     c.drawString(100, y, f'{key}: {value}')
    #     y -= 20
    for key, value in list(defects.items())[:-1]:
        c.drawString(100, y, f'{key}: {value}')
        y -= 20
    c.save()

    with open(file_path, 'rb') as f:
        response = JsonResponse({'status': 'success'})
        response['Content-Disposition'] = f'attachment; filename="report_{serial_number}.pdf"'
        response.write(f.read())
        return response






